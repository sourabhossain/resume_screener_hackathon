"""
Document text extraction from PDF and DOCX files.
"""
import logging
from pathlib import Path

from pypdf import PdfReader
import docx

logger = logging.getLogger(__name__)


# Fancy resume templates render icons/labels using symbol fonts, which PDF
# extraction turns into decorative Unicode glyphs (e.g. ⌢ U+2322, bullets,
# en/em dashes, smart quotes). The raw_text column is utf8mb4 so these store
# fine, but we still normalise the common readable ones to ASCII so the LLM
# sees clean text rather than icon-font artefacts.
_CHAR_REPLACEMENTS = {
    '–': '-', '—': '-', '−': '-',
    '•': '*', '·': '*', '●': '*', '○': '*',
    '“': '"', '”': '"', '‘': "'", '’': "'",
    '…': '...', ' ': ' ',
}


class DocumentExtractor:
    """Extract text content from resume documents."""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}

    @classmethod
    def extract(cls, file_path: str) -> str:
        """
        Extract text from a document file.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension == '.pdf':
            text = cls._extract_from_pdf(file_path)
        elif extension == '.docx':
            text = cls._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        return cls._sanitize(text)

    @staticmethod
    def _sanitize(text: str) -> str:
        """Replace common typographic glyphs with ASCII equivalents.

        The DB column is utf8mb4 (inherits the table charset), so all Unicode
        is stored correctly. We only normalise decorative symbols so the LLM
        sees clean text rather than icon-font artefacts.
        """
        if not text:
            return text

        for src, dst in _CHAR_REPLACEMENTS.items():
            text = text.replace(src, dst)

        return text
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF, preferring PyMuPDF.

        pypdf renders some styled/letter-spaced templates character-by-character
        (a name heading comes out as 'M a s h r u r  M a h m u d'), which breaks
        name/skill extraction. PyMuPDF reconstructs words correctly on those same
        files and matches pypdf on plain PDFs, so it is the primary extractor;
        pypdf stays as a fallback if PyMuPDF is unavailable or errors.
        """
        text = DocumentExtractor._extract_pdf_pymupdf(file_path)
        if text and text.strip():
            return text
        return DocumentExtractor._extract_pdf_pypdf(file_path)

    @staticmethod
    def _extract_pdf_pymupdf(file_path: str) -> str:
        """Primary PDF extractor (PyMuPDF). Returns '' on any failure so the
        caller can fall back to pypdf.

        Imported as `pymupdf`, its current name. The old `fitz` alias still works
        but warns on every call and is slated for removal, so it is only a
        fallback for an older wheel.
        """
        try:
            import pymupdf
        except ImportError:
            try:
                import fitz as pymupdf  # PyMuPDF < 1.24.3
            except ImportError:
                logger.warning("PyMuPDF not installed; falling back to pypdf")
                return ''
        try:
            doc = pymupdf.open(file_path)
            try:
                return "\n".join(page.get_text() for page in doc)
            finally:
                doc.close()
        except Exception as e:
            logger.warning("PyMuPDF extraction failed for %s (%s); falling back to pypdf", file_path, e)
            return ''

    @staticmethod
    def _extract_pdf_pypdf(file_path: str) -> str:
        """Fallback PDF extractor (pypdf)."""
        text_parts = []

        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return "\n".join(text_parts)
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        text_parts = []
        

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if file type is supported."""
        extension = Path(file_path).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS
